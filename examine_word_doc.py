from docx import Document
import os

def examine_word_document(docx_path):
    """Examine the structure of the Word document"""
    try:
        doc = Document(docx_path)
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        print("\n" + "="*50)
        print("DOCUMENT STRUCTURE:")
        print("="*50)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"Paragraph {i+1}:")
                print(f"  Text: {text[:100]}{'...' if len(text) > 100 else ''}")
                print(f"  Length: {len(text)}")
                print(f"  Style: {paragraph.style.name if paragraph.style else 'No style'}")
                
                # Check if paragraph has runs with formatting
                if paragraph.runs:
                    print(f"  Runs: {len(paragraph.runs)}")
                    for j, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            print(f"    Run {j+1}: '{run.text[:50]}{'...' if len(run.text) > 50 else ''}' (Bold: {run.bold}, Italic: {run.italic})")
                print("-" * 30)
        
        # Also check for tables
        if doc.tables:
            print(f"\nDocument has {len(doc.tables)} tables")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    print(f"  Row {row_idx+1}: {' | '.join(row_text)}")
        
    except Exception as e:
        print(f"Error examining document: {e}")

if __name__ == "__main__":
    examine_word_document("My Projects.docx") 